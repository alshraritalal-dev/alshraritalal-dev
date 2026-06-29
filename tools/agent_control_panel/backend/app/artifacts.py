from __future__ import annotations

import ast
import json
import mimetypes
import shutil
import zipfile
from io import BytesIO
from pathlib import Path
from typing import Any

from PIL import Image

try:
    from docx import Document
except Exception:  # pragma: no cover - optional dependency guard
    Document = None

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - optional dependency guard
    PdfReader = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
except Exception:  # pragma: no cover - optional dependency guard
    A4 = None
    canvas = None

from .config import Settings
from .helpers import utc_now_iso
from .schemas import RunArtifact


TEXT_SUFFIXES = {
    ".py",
    ".md",
    ".json",
    ".html",
    ".txt",
    ".csv",
    ".css",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".vue",
    ".yaml",
    ".yml",
    ".toml",
    ".xml",
    ".log",
    ".ini",
}

DOCUMENT_SUFFIXES = {".pdf", ".docx"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}


class ArtifactService:
    def __init__(self, settings: Settings) -> None:
        self.settings = settings
        self.root = settings.artifact_root
        self.root.mkdir(parents=True, exist_ok=True)

    def run_directory(self, run_id: str) -> Path:
        target = self.root / run_id
        target.mkdir(parents=True, exist_ok=True)
        return target

    def write_content_to_path(
        self,
        target: Path,
        content: str,
        *,
        encoding: str = "utf-8",
        append: bool = False,
    ) -> dict[str, Any]:
        suffix = target.suffix.lower()
        target.parent.mkdir(parents=True, exist_ok=True)

        if append and suffix in DOCUMENT_SUFFIXES:
            existing = self.extract_text_for_prompt(target) if target.exists() else ""
            content = f"{existing.rstrip()}\n{content}".strip()
            append = False

        if suffix == ".pdf":
            self._write_pdf(target, content)
        elif suffix == ".docx":
            self._write_docx(target, content)
        elif suffix == ".json":
            normalized = self._normalize_json_text(content)
            self._write_text(target, normalized, encoding=encoding, append=append)
        else:
            self._write_text(target, content, encoding=encoding, append=append)

        return self.verify_output(target)

    def stage_file(
        self,
        run_id: str,
        source_path: Path,
        *,
        generated_by_agent: str,
        workspace_root: str | None = None,
        description: str = "",
    ) -> RunArtifact:
        resolved_source = source_path.expanduser().resolve()
        run_dir = self.run_directory(run_id)
        destination = self._stage_destination(run_dir, resolved_source, workspace_root)

        if resolved_source != destination:
            destination.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(resolved_source, destination)

        metadata = self._build_artifact_metadata(
            run_id,
            destination,
            generated_by_agent=generated_by_agent,
            description=description,
            source_absolute_path=str(resolved_source),
        )
        return metadata

    def bundle_run_artifacts(self, run_id: str, file_name: str = "artifacts_bundle.zip") -> RunArtifact:
        run_dir = self.run_directory(run_id)
        zip_path = run_dir / file_name
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for path in sorted(run_dir.rglob("*")):
                if not path.is_file() or path == zip_path:
                    continue
                archive.write(path, arcname=path.relative_to(run_dir))
        return self._build_artifact_metadata(run_id, zip_path, generated_by_agent="system")

    def extract_text_for_prompt(self, path: Path | str, max_chars: int = 8000) -> str:
        target = Path(path).expanduser().resolve()
        suffix = target.suffix.lower()

        if suffix in TEXT_SUFFIXES:
            return target.read_text(encoding="utf-8", errors="replace")[:max_chars]

        if suffix == ".pdf":
            if PdfReader is None:
                raise RuntimeError("pypdf is required to read PDF files.")
            reader = PdfReader(str(target))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text[:max_chars]

        if suffix == ".docx":
            if Document is None:
                raise RuntimeError("python-docx is required to read DOCX files.")
            doc = Document(str(target))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            return text[:max_chars]

        return ""

    def verify_output(self, target: Path) -> dict[str, Any]:
        resolved = target.expanduser().resolve()
        if not resolved.exists():
            raise RuntimeError(f"Output file does not exist after write: {resolved}")

        size = resolved.stat().st_size
        if size <= 0:
            raise RuntimeError(f"Output file is empty: {resolved}")

        suffix = resolved.suffix.lower()

        if suffix == ".py":
            ast.parse(resolved.read_text(encoding="utf-8", errors="strict"))
        elif suffix == ".json":
            json.loads(resolved.read_text(encoding="utf-8", errors="strict"))
        elif suffix == ".pdf":
            if PdfReader is None:
                raise RuntimeError("pypdf is required to validate PDF files.")
            reader = PdfReader(str(resolved))
            if len(reader.pages) < 1:
                raise RuntimeError(f"PDF has no readable pages: {resolved}")
        elif suffix == ".docx":
            if Document is None:
                raise RuntimeError("python-docx is required to validate DOCX files.")
            doc = Document(str(resolved))
            if not doc.paragraphs and size < 512:
                raise RuntimeError(f"DOCX content looks invalid: {resolved}")
        elif suffix in IMAGE_SUFFIXES:
            with Image.open(resolved) as image:
                image.verify()
        elif suffix in TEXT_SUFFIXES:
            if not resolved.read_text(encoding="utf-8", errors="strict").strip():
                raise RuntimeError(f"Text output is blank: {resolved}")

        return {
            "ok": True,
            "path": str(resolved),
            "size_bytes": size,
            "extension": suffix,
        }

    def _write_text(self, target: Path, content: str, *, encoding: str, append: bool) -> None:
        if append:
            with target.open("a", encoding=encoding, errors="strict", newline="") as handle:
                handle.write(content)
            return
        target.write_text(content, encoding=encoding, errors="strict")

    def _write_pdf(self, target: Path, content: str) -> None:
        if canvas is None or A4 is None:
            raise RuntimeError("reportlab is required to generate PDF files.")

        buffer = BytesIO()
        pdf = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        margin = 48
        text_object = pdf.beginText(margin, height - margin)
        text_object.setFont("Helvetica", 11)
        lines = (content or "").splitlines() or [""]

        for raw_line in lines:
            line = raw_line.rstrip()
            segments = self._wrap_pdf_line(line, max_chars=95) or [""]
            for segment in segments:
                if text_object.getY() <= margin:
                    pdf.drawText(text_object)
                    pdf.showPage()
                    text_object = pdf.beginText(margin, height - margin)
                    text_object.setFont("Helvetica", 11)
                text_object.textLine(segment)

        pdf.drawText(text_object)
        pdf.save()
        target.write_bytes(buffer.getvalue())

    def _write_docx(self, target: Path, content: str) -> None:
        if Document is None:
            raise RuntimeError("python-docx is required to generate DOCX files.")
        doc = Document()
        for block in (content or "").splitlines() or [""]:
            doc.add_paragraph(block)
        doc.save(str(target))

    def _normalize_json_text(self, content: str) -> str:
        try:
            data = json.loads(content)
        except json.JSONDecodeError:
            return content
        return json.dumps(data, ensure_ascii=False, indent=2) + "\n"

    def _stage_destination(self, run_dir: Path, source: Path, workspace_root: str | None) -> Path:
        if self._is_within(source, run_dir):
            return source

        segments: list[str] = []
        if workspace_root:
            workspace_path = Path(workspace_root).expanduser().resolve()
            if self._is_within(source, workspace_path):
                relative = source.relative_to(workspace_path)
                segments = ["workspace", *relative.parts]

        if not segments:
            segments = [source.name]

        destination = run_dir.joinpath(*segments)
        return destination

    def _build_artifact_metadata(
        self,
        run_id: str,
        staged_path: Path,
        *,
        generated_by_agent: str,
        description: str = "",
        source_absolute_path: str | None = None,
    ) -> RunArtifact:
        resolved = staged_path.expanduser().resolve()
        mime_type = mimetypes.guess_type(resolved.name)[0] or "application/octet-stream"
        preview_kind = self._preview_kind(resolved, mime_type)
        preview_text = self._preview_text(resolved, preview_kind)
        artifact = RunArtifact(
            absolute_path=str(resolved),
            source_absolute_path=source_absolute_path,
            file_name=resolved.name,
            file_extension=resolved.suffix.lower().lstrip("."),
            mime_type=mime_type,
            file_size=resolved.stat().st_size,
            generated_by_agent=generated_by_agent,
            previewable=preview_kind != "binary",
            preview_kind=preview_kind,
            preview_text=preview_text,
            description=description,
            open_url=f"/api/artifacts/open?run_id={run_id}&artifact_id={{artifact_id}}",
            download_url=f"/api/artifacts/download?run_id={run_id}&artifact_id={{artifact_id}}",
            created_at=utc_now_iso(),
        )
        artifact.open_url = artifact.open_url.format(artifact_id=artifact.id)
        artifact.download_url = artifact.download_url.format(artifact_id=artifact.id)
        return artifact

    def _preview_kind(self, path: Path, mime_type: str) -> str:
        suffix = path.suffix.lower()
        if suffix in IMAGE_SUFFIXES or mime_type.startswith("image/"):
            return "image"
        if suffix in TEXT_SUFFIXES or mime_type.startswith("text/"):
            return "text"
        if suffix in DOCUMENT_SUFFIXES:
            return "document"
        return "binary"

    def _preview_text(self, path: Path, preview_kind: str) -> str:
        try:
            if preview_kind == "text":
                return path.read_text(encoding="utf-8", errors="replace")[:1600]
            if preview_kind == "document":
                return self.extract_text_for_prompt(path, max_chars=1600)
        except Exception:
            return ""
        return ""

    def _wrap_pdf_line(self, line: str, max_chars: int) -> list[str]:
        if len(line) <= max_chars:
            return [line]

        parts: list[str] = []
        current = line
        while len(current) > max_chars:
            parts.append(current[:max_chars])
            current = current[max_chars:]
        if current:
            parts.append(current)
        return parts

    def _is_within(self, candidate: Path, root: Path) -> bool:
        try:
            candidate.resolve().relative_to(root.resolve())
            return True
        except ValueError:
            return False
